import os
import fitz  # PyMuPDF
from langchain_community.document_loaders.pdf import PyPDFLoader
from docx import Document as DocxDocument
from fastapi import FastAPI, HTTPException, Form
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain_mistralai.embeddings import MistralAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import create_retrieval_chain
from langchain.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_mistralai.chat_models import ChatMistralAI
import io
import json
from typing import List
from langchain.docstore.document import Document
from mistralai.client import MistralClient
from mistralai.models.chat_completion import ChatMessage

app = FastAPI()

class Query(BaseModel):
    query: str

class QueryRequest(BaseModel):
    query: str

folder_id = None
client = MistralClient(api_key="_")

# Global variables to cache documents and vector store
cached_documents = []
cached_vector_store = None

# Path to cache file
CACHE_FILE = 'document_cache.json'

def save_cache(documents):
    with open(CACHE_FILE, 'w') as f:
        json.dump([doc.dict() for doc in documents], f)

def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, 'r') as f:
            return [Document(**doc) for doc in json.load(f)]
    return []

@app.post("/set_folder/")
async def set_folder(folder_id_input: str = Form(...)):
    global folder_id, cached_documents, cached_vector_store
    folder_id = folder_id_input

    SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
    SERVICE_ACCOUNT_FILE = 'Path/to/service/account/json/file'  # Update this path

    credentials = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    service = build('drive', 'v3', credentials=credentials)

    def list_files_in_folder(service, folder_id):
        query = f"'{folder_id}' in parents and trashed=false"
        try:
            results = service.files().list(q=query, pageSize=10, fields="files(id, name)").execute()
            items = results.get('files', [])
            if not items:
                print(f"No files found in the folder with ID: {folder_id}")
            return items
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    files = list_files_in_folder(service, folder_id)
    if not files:
        raise ValueError(f"No files found or unable to access the folder with ID: {folder_id}")

    documents = []
    for file in files:
        file_id = file['id']
        file_name = file['name']
        print(f"Processing file: {file_name} (ID: {file_id})")
        request = service.files().get_media(fileId=file_id)
        file_io = io.BytesIO()
        downloader = MediaIoBaseDownload(file_io, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        file_io.seek(0)

        if file_name.endswith('.pdf'):
            text = extract_text_from_pdf(file_io)
        elif file_name.endswith('.docx'):
            text = extract_text_from_docx(file_io)
        else:
            text = file_io.read().decode('utf-8')

        documents.append(Document(page_content=text))

    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter()
    split_documents = text_splitter.split_documents(documents)

    os.environ["HF_TOKEN"] = "your_huggingface_token" 

    mistral_api_key = "_"
    embeddings = MistralAIEmbeddings(model="mistral-embed", api_key=mistral_api_key)

    if not split_documents:
        raise HTTPException(status_code=500, detail="No documents found to process")

    cached_vector_store = FAISS.from_documents(split_documents, embeddings)
    cached_documents = split_documents

    # Save the cache
    save_cache(cached_documents)

    return {"message": "Folder ID set and documents indexed successfully"}

def extract_text_from_pdf(file_io):
    text = ""
    document = fitz.open(stream=file_io, filetype="pdf")
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        text += page.get_text()
    return text

def extract_text_from_docx(file_io):
    doc = DocxDocument(file_io)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

@app.post("/query")
async def query_drive(query_request: QueryRequest):
    global cached_vector_store
    if not folder_id:
        raise HTTPException(status_code=400, detail="Folder ID not set")

    if not cached_documents:
        raise HTTPException(status_code=500, detail="Documents are not loaded. Please set the folder ID again.")

    retriever = cached_vector_store.as_retriever()

    mistral_api_key = "_"
    model = ChatMistralAI(api_key=mistral_api_key)

    prompt = ChatPromptTemplate.from_template("""Answer the following question based only on the provided context:

    <context>
    {context}
    </context>

    Question: {input}""")

    document_chain = create_stuff_documents_chain(model, prompt)
    retrieval_chain = create_retrieval_chain(retriever, document_chain)

    try:
        response = retrieval_chain.invoke({"input": query_request.query})
        return {"answer": response["answer"]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/", response_class=HTMLResponse)
async def read_root():
    with open("index.html", "r") as file:
        return HTMLResponse(content=file.read(), status_code=200)
    
if __name__ == "__main__":
    import threading
    import uvicorn

    def run_app():
        uvicorn.run(app, host="0.0.0.0", port=8000)

    thread = threading.Thread(target=run_app, args=(), daemon=True)
    thread.start()
